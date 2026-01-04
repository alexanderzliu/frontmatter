"""Word document parser."""

import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from typeset.config.models import StyleMapping
from typeset.ir.document import Chapter, Document, Metadata
from typeset.ir.nodes import (
    FootnoteNode,
    FootnoteRefNode,
    HeadingNode,
    ImageNode,
    ListNode,
    Node,
    NodeType,
    TableNode,
    TextNode,
    text,
)


class DocxParser:
    """Parse Word documents into IR."""

    def __init__(self, style_mapping: Optional[StyleMapping] = None):
        self.style_mapping = style_mapping or StyleMapping()
        self.footnote_counter = 0
        self.footnotes: dict[str, Node] = {}
        self.images: dict[str, bytes] = {}
        self.image_counter = 0

    def parse(self, file_path: str | Path) -> Document:
        """Parse a .docx file into IR Document."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        docx = DocxDocument(str(file_path))

        metadata = self._extract_metadata(docx, file_path)
        self._extract_images(docx)
        chapters = self._extract_chapters(docx)

        return Document(
            metadata=metadata,
            chapters=chapters,
            footnotes=self.footnotes,
            images=self.images,
        )

    def _extract_metadata(self, docx: DocxDocument, file_path: Path) -> Metadata:
        """Extract document metadata from core properties."""
        props = docx.core_properties

        return Metadata(
            title=props.title or file_path.stem,
            authors=[props.author] if props.author else [],
            language=props.language or "en",
            description=props.comments or None,
            keywords=props.keywords.split(",") if props.keywords else [],
        )

    def _extract_images(self, docx: DocxDocument) -> None:
        """Extract embedded images from document."""
        for rel in docx.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    self.images[rel.rId] = image_data
                except Exception:
                    pass  # Skip problematic images

    def _extract_chapters(self, docx: DocxDocument) -> list[Chapter]:
        """Split document into chapters based on heading styles."""
        chapters: list[Chapter] = []
        current_chapter: Optional[Chapter] = None
        chapter_counter = 0

        for element in docx.element.body:
            # Handle paragraphs
            if element.tag.endswith("p"):
                para = Paragraph(element, docx)
                style_name = para.style.name if para.style else "Normal"

                if self._is_chapter_heading(style_name):
                    # Save current chapter
                    if current_chapter:
                        chapters.append(current_chapter)

                    # Start new chapter
                    chapter_counter += 1
                    current_chapter = Chapter(
                        title=para.text.strip() or f"Chapter {chapter_counter}",
                        level=self._heading_level(style_name),
                        content=[],
                        id=f"chapter-{chapter_counter}",
                    )
                elif current_chapter:
                    # Parse and add to current chapter
                    if self._is_section_heading(style_name):
                        node = self._parse_heading(para)
                    else:
                        node = self._parse_paragraph(para)
                    if node:
                        current_chapter.content.append(node)
                else:
                    # Content before first chapter heading
                    node = self._parse_paragraph(para)
                    if node and node.get_text().strip():
                        # Create implicit first chapter
                        chapter_counter += 1
                        current_chapter = Chapter(
                            title="",
                            level=1,
                            content=[node],
                            id=f"chapter-{chapter_counter}",
                        )

            # Handle tables
            elif element.tag.endswith("tbl"):
                table = Table(element, docx)
                node = self._parse_table(table)
                if current_chapter and node:
                    current_chapter.content.append(node)

        # Don't forget the last chapter
        if current_chapter:
            chapters.append(current_chapter)

        # If no chapters were created, create one with all content
        if not chapters:
            chapters = [
                Chapter(
                    title="",
                    level=1,
                    content=[],
                    id="chapter-1",
                )
            ]

        return chapters

    def _is_chapter_heading(self, style_name: str) -> bool:
        """Check if style indicates a chapter heading."""
        style_lower = style_name.lower()
        for pattern in self.style_mapping.chapter_heading_styles:
            if pattern.lower() in style_lower or style_lower in pattern.lower():
                return True
        return False

    def _is_section_heading(self, style_name: str) -> bool:
        """Check if style indicates a section heading."""
        style_lower = style_name.lower()
        for pattern in self.style_mapping.section_heading_styles:
            if pattern.lower() in style_lower or style_lower in pattern.lower():
                return True
        return False

    def _heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        match = re.search(r"(\d+)", style_name)
        if match:
            level = int(match.group(1))
            return min(max(level, 1), 6)  # Clamp to 1-6
        return 1

    def _parse_heading(self, para: Paragraph) -> HeadingNode:
        """Parse a heading paragraph."""
        style_name = para.style.name if para.style else "Heading 2"
        level = self._heading_level(style_name)

        children = self._parse_runs(para.runs)
        return HeadingNode(
            node_type=NodeType.HEADING, level=level, children=children if children else [text(para.text)]
        )

    def _parse_paragraph(self, para: Paragraph) -> Optional[Node]:
        """Parse a paragraph into an IR node."""
        # Skip empty paragraphs
        if not para.text.strip() and not self._has_images(para):
            return None

        # Check for images
        images = self._extract_paragraph_images(para)
        if images and not para.text.strip():
            # Paragraph with only images
            if len(images) == 1:
                return images[0]
            return Node(node_type=NodeType.PARAGRAPH, children=images)

        # Parse runs (text with formatting)
        children = self._parse_runs(para.runs)

        # Add any images
        children.extend(images)

        if not children:
            return None

        # Check for blockquote style
        style_name = para.style.name if para.style else "Normal"
        if any(bq.lower() in style_name.lower() for bq in self.style_mapping.blockquote_styles):
            return Node(node_type=NodeType.BLOCKQUOTE, children=children)

        return Node(node_type=NodeType.PARAGRAPH, children=children)

    def _parse_runs(self, runs: list[Run]) -> list[Node]:
        """Parse a list of runs into nodes."""
        nodes: list[Node] = []

        for run in runs:
            if not run.text:
                continue

            node: Node = TextNode(node_type=NodeType.TEXT, text=run.text)

            # Apply formatting (innermost first)
            if run.italic:
                node = Node(node_type=NodeType.EMPHASIS, children=[node])
            if run.bold:
                node = Node(node_type=NodeType.STRONG, children=[node])
            if run.font.strike:
                node = Node(node_type=NodeType.STRIKETHROUGH, children=[node])
            if run.font.superscript:
                node = Node(node_type=NodeType.SUPERSCRIPT, children=[node])
            if run.font.subscript:
                node = Node(node_type=NodeType.SUBSCRIPT, children=[node])

            nodes.append(node)

        return nodes

    def _has_images(self, para: Paragraph) -> bool:
        """Check if paragraph contains images."""
        # Use lxml nsmap for namespace-aware xpath
        nsmap = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        try:
            return bool(para._element.xpath(".//a:blip", namespaces=nsmap))
        except TypeError:
            # Fallback: search without namespace
            return bool(para._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))

    def _extract_paragraph_images(self, para: Paragraph) -> list[ImageNode]:
        """Extract images from a paragraph."""
        images: list[ImageNode] = []

        # Find all blip elements (embedded images)
        nsmap = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        try:
            blips = para._element.xpath(".//a:blip", namespaces=nsmap)
        except TypeError:
            # Fallback: use findall with full namespace
            blips = para._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")

        for blip in blips:
            embed_id = blip.get(qn("r:embed"))
            if embed_id and embed_id in self.images:
                self.image_counter += 1
                image_data = self.images[embed_id]

                # Determine mime type from data
                mime_type = self._detect_image_type(image_data)

                images.append(
                    ImageNode(
                        node_type=NodeType.IMAGE,
                        src=image_data,
                        mime_type=mime_type,
                        filename=f"image_{self.image_counter}.{mime_type.split('/')[-1]}",
                    )
                )

        return images

    def _detect_image_type(self, data: bytes) -> str:
        """Detect image MIME type from data."""
        if data.startswith(b"\x89PNG"):
            return "image/png"
        elif data.startswith(b"\xff\xd8\xff"):
            return "image/jpeg"
        elif data.startswith(b"GIF"):
            return "image/gif"
        elif data.startswith(b"RIFF") and b"WEBP" in data[:12]:
            return "image/webp"
        return "image/png"  # Default

    def _parse_table(self, table: Table) -> TableNode:
        """Parse a table into an IR node."""
        rows: list[Node] = []

        for row in table.rows:
            cells: list[Node] = []
            for cell in row.cells:
                # Parse cell content
                cell_content: list[Node] = []
                for para in cell.paragraphs:
                    node = self._parse_paragraph(para)
                    if node:
                        cell_content.append(node)

                cell_node = Node(node_type=NodeType.TABLE_CELL, children=cell_content)
                cells.append(cell_node)

            row_node = Node(node_type=NodeType.TABLE_ROW, children=cells)
            rows.append(row_node)

        return TableNode(node_type=NodeType.TABLE, children=rows)
